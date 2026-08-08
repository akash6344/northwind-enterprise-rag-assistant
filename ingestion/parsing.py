from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document

from ingestion.metadata import extract_common_metadata


def parse_pdf(path: Path) -> list[dict]:
    records = []
    with pdfplumber.open(path) as pdf:
        full_text = "\n".join((page.extract_text() or "") for page in pdf.pages)
        base = extract_common_metadata(path, full_text)
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                records.append({**base, "page": idx, "text": text})
    return records


def parse_docx(path: Path) -> list[dict]:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table_idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
        if rows:
            header = rows[0]
            lines = [f"Table {table_idx}:"]
            for row in rows[1:]:
                lines.append(" | ".join(f"{header[i] if i < len(header) else 'Column'}: {value}" for i, value in enumerate(row)))
            parts.append("\n".join(lines))
    text = "\n".join(parts)
    base = extract_common_metadata(path, text)
    return [{**base, "page": None, "text": text}]


def parse_xlsx(path: Path) -> list[dict]:
    xl = pd.ExcelFile(path)
    records = []
    preview_text = ""
    for sheet in xl.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet, header=None)
        df = df.dropna(how="all").dropna(axis=1, how="all")
        markdown = dataframe_to_text_table(df)
        text = f"Worksheet: {sheet}\n{markdown}"
        preview_text += "\n" + text
        records.append({"sheet": sheet, "page": None, "text": text})
    base = extract_common_metadata(path, preview_text)
    return [{**base, **record} for record in records]


def dataframe_to_text_table(df: pd.DataFrame) -> str:
    lines = []
    for row in df.astype(object).where(pd.notna(df), "").values.tolist():
        values = [str(value).strip() for value in row]
        if any(values):
            lines.append(" | ".join(values))
    return "\n".join(lines)


def parse_file(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".xlsx":
        return parse_xlsx(path)
    return []


def parse_knowledge_base(root: Path) -> list[dict]:
    records: list[dict] = []
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.name != ".DS_Store":
            records.extend(parse_file(path))
    return records
